"""
P15-S11 D2 — text extraction unit tests.

Coverage focus:
  * PDF text-layer happy path (uses VinFast PO sample — ships in repo).
  * DOCX synthetic happy path (built inline).
  * Image short-circuit (PNG bytes → 'unsupported_today').
  * Unknown mime → 'failed'.
  * Empty PDF → 'unsupported_today' (scanned-PDF surrogate).
  * Encrypted PDF → 'failed' with friendly Vietnamese hint.
  * Cache fingerprint stable across calls.

Pure Python — no DB, no FastAPI, no Postgres.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from data_pipeline.data_plane.silver.docsage_extract import (
    ExtractResult,
    content_fingerprint,
    extract_text,
)


REPO = Path(__file__).resolve().parent.parent.parent.parent
VINFAST_PDF_DIR = REPO / "data" / "vinfast" / "vinfast_test_bundle" / "sample_documents"


# ─── PDF — happy path on real VinFast sample ────────────────────────


class TestPdfTextExtraction:

    @pytest.fixture
    def vinfast_po_bytes(self) -> bytes:
        path = VINFAST_PDF_DIR / "PO_VF-WO-26057_DealerPurchaseOrder.pdf"
        if not path.exists():
            pytest.skip(f"VinFast PDF fixture not found at {path}")
        return path.read_bytes()

    def test_happy_path_returns_ok_status(self, vinfast_po_bytes):
        r = extract_text(content=vinfast_po_bytes, mime_type="application/pdf")
        assert r.status == "ok", f"Expected ok, got {r.status!r}: {r.error_message}"
        assert r.page_count >= 1
        assert r.char_count > 0
        assert len(r.text) == r.char_count

    def test_page_offsets_align_with_text(self, vinfast_po_bytes):
        r = extract_text(content=vinfast_po_bytes, mime_type="application/pdf")
        # page_offsets has len = page_count + 1 so substring slicing works:
        # text[offsets[i]:offsets[i+1]] is page i+1's text (plus the \n
        # separator we add).
        assert len(r.page_offsets) == r.page_count + 1
        # The last offset must be ≥ len(text) (we add 1 per page for the
        # separator).
        assert r.page_offsets[-1] >= len(r.text)

    def test_filename_fallback_when_mime_missing(self, vinfast_po_bytes):
        r = extract_text(
            content=vinfast_po_bytes, mime_type="application/octet-stream",
            filename="purchase_order.pdf",
        )
        assert r.status == "ok"

    def test_empty_pdf_marked_unsupported_today(self):
        """A minimal valid PDF with no extractable text — surrogate for
        a scanned-PDF input. extract_text() returns empty per page and
        we mark the whole result `unsupported_today` (FE sends to OCR
        queue when the Phase 2 adapter lands)."""
        # Smallest valid PDF: 1 empty page
        empty_pdf = (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj\n"
            b"xref\n0 4\n0000000000 65535 f\n0000000009 00000 n\n"
            b"0000000052 00000 n\n0000000096 00000 n\n"
            b"trailer<</Size 4/Root 1 0 R>>\nstartxref\n149\n%%EOF\n"
        )
        r = extract_text(content=empty_pdf, mime_type="application/pdf")
        assert r.status == "unsupported_today"
        assert "scan" in (r.error_message or "").lower() or "OCR" in (r.error_message or "")
        assert r.page_count == 1

    def test_corrupt_pdf_returns_failed(self):
        r = extract_text(content=b"not a pdf at all", mime_type="application/pdf")
        assert r.status == "failed"
        assert "không đọc được" in r.error_message


# ─── DOCX — built inline ────────────────────────────────────────────


class TestDocxTextExtraction:

    def _build_docx(self, *paragraphs: str, table: list[list[str]] | None = None) -> bytes:
        doc = DocxDocument()
        for p in paragraphs:
            doc.add_paragraph(p)
        if table:
            t = doc.add_table(rows=len(table), cols=len(table[0]))
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    t.rows[i].cells[j].text = cell
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_happy_path_paragraphs_only(self):
        b = self._build_docx("Hợp đồng số HD-2026-001", "Bên A: Vingroup")
        r = extract_text(content=b, mime_type=
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert r.status == "ok"
        assert "HD-2026-001" in r.text
        assert "Vingroup" in r.text
        assert r.page_count == 1

    def test_table_cells_included(self):
        """Financial DOCX usually carries a table — we MUST include
        table cells in the extracted text or DocSage misses the
        most important rows."""
        b = self._build_docx(
            "Báo cáo Quý 1",
            table=[
                ["Chỉ tiêu", "Doanh thu", "Lợi nhuận"],
                ["Q1/2026",  "100M",      "20M"],
            ],
        )
        r = extract_text(content=b, mime_type=
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert r.status == "ok"
        assert "Doanh thu" in r.text
        assert "Q1/2026" in r.text

    def test_empty_docx_returns_failed(self):
        b = self._build_docx()
        r = extract_text(content=b, mime_type=
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert r.status == "failed"

    def test_filename_fallback(self):
        b = self._build_docx("Một dòng")
        r = extract_text(
            content=b, mime_type="application/octet-stream",
            filename="contract.docx",
        )
        assert r.status == "ok"


# ─── Image short-circuit ────────────────────────────────────────────


class TestImageShortCircuit:

    def test_png_marked_unsupported_today(self):
        r = extract_text(
            content=b"\x89PNG\r\n\x1a\n",  # just the magic bytes
            mime_type="image/png",
        )
        assert r.status == "unsupported_today"
        assert "Qwen2-VL" in r.error_message or "Phase 2" in r.error_message

    def test_jpg_marked_unsupported_today(self):
        r = extract_text(content=b"\xff\xd8\xff", mime_type="image/jpeg")
        assert r.status == "unsupported_today"


# ─── Unknown mime ───────────────────────────────────────────────────


class TestUnknownFormat:

    def test_unknown_mime_returns_failed(self):
        r = extract_text(content=b"some bytes", mime_type="application/x-foo",
                          filename="thing.foo")
        assert r.status == "failed"
        assert "Định dạng không hỗ trợ" in r.error_message


# ─── Fingerprint cache key ──────────────────────────────────────────


class TestContentFingerprint:

    def test_same_bytes_same_fingerprint(self):
        b = b"hello world" * 100
        a = content_fingerprint(b, lib_version="pypdf-5.0.1")
        c = content_fingerprint(b, lib_version="pypdf-5.0.1")
        assert a == c

    def test_lib_version_changes_fingerprint(self):
        b = b"hello world" * 100
        a = content_fingerprint(b, lib_version="pypdf-5.0.1")
        c = content_fingerprint(b, lib_version="pypdf-5.1.0")
        assert a != c, "Lib version must invalidate cache."

    def test_different_bytes_different_fingerprint(self):
        a = content_fingerprint(b"alpha", lib_version="v1")
        c = content_fingerprint(b"beta", lib_version="v1")
        assert a != c
