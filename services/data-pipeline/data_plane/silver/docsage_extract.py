"""
Stage 6 — Text extraction for unstructured docs (PDF / DOCX).

This is the Silver-tier transform that turns the placeholder
`unstructured_pending` state (shipped 2026-05-17 commit 8494608) into
actual queryable text. The downstream consumer is DocSage (P15-S11
D3-D6): Schema Discovery reads the extracted text out of
`bronze_files.metadata.docsage_text`; Structured Extraction reads the
per-page boundaries to emit `source_segment` citations.

Per Medallion separation (see feedback_medallion_separation): this
module lives in `silver/`, NOT `bronze/`. Bronze keeps the raw bytes;
Silver carries the parsed-and-typed view. The extraction is a Silver
transform of an unstructured Bronze file.

OCR (scanned PDFs / images) is **deferred to Phase 2** — we need the
Qwen2-VL vision-LLM adapter via llm-gateway to do this without
shipping yet-another inference dep. For now image-only PDFs and image
files are flagged `extraction_status='unsupported_today'` and the
caller surfaces a friendly Vietnamese explanation on the FE.

K-3 N/A here — extraction is deterministic Python lib calls, no LLM.
K-5 N/A here too — text stays inside the tenant; no external API
involved at this layer. DocSage layers above apply K-5 before any
external-vendor call (per D4 in the plan).

Side_effect_class = `write_idempotent` — same bytes + same lib version
= same output.
"""
from __future__ import annotations

import hashlib
import io
from dataclasses import dataclass, field
from typing import Optional

import pypdf
from docx import Document as DocxDocument

from .blocks import Bbox, Block, BlockType
from .header_footer_strip import strip_repeating_lines
from .reading_order import (
    extract_reading_order_from_pdf,
    multi_column_page_count,
)
from .table_extractor import extract_tables_from_pdf


# ─── Shape ──────────────────────────────────────────────────────────


@dataclass(frozen=True)
class ExtractResult:
    """Output of the text-extraction step. The 4 fields tell DocSage
    everything it needs:

      * `text` — the plain UTF-8 to feed prompts.
      * `page_offsets` — page_index → start_char_offset in `text`. A
        DocSage extraction Row carrying `source_segment=(3, 5)` becomes
        the substring `text[page_offsets[3]:page_offsets[6]]` so the FE
        can highlight the right pages on the original PDF.
      * `status` — one of {'ok', 'partial', 'unsupported_today',
        'failed'}. 'partial' means we got page text from some pages
        but errors on others (e.g. corrupted PDF objects).
      * `error_message` — free-text reason for `partial` / 'failed' /
        'unsupported_today'. Rendered in Vietnamese on the FE.
    """
    text:          str
    page_offsets:  list[int]       # len = num_pages + 1; offsets into `text`
    status:        str             # 'ok' / 'partial' / 'unsupported_today' / 'failed'
    error_message: Optional[str] = None
    page_count:    int             = 0
    char_count:    int             = 0
    # MinerU pattern 1+2 (P2-05-18) — optional block taxonomy + header
    # /footer stripping. Old callers ignore `blocks`; new callers (FE
    # bbox highlight, DocSage Schema Discovery v2, Pattern 3 tables)
    # read it for structural signal.
    blocks:        Optional[list[Block]] = field(default=None)
    header_footer_lines_stripped: int = 0
    # Pattern 4 — count of pages where multi-column reading-order
    # reconstruction kicked in. 0 = every page was single-column or
    # pdfplumber unavailable.
    multi_column_pages_reordered: int = 0


# ─── Public API ─────────────────────────────────────────────────────


def extract_text(
    *, content: bytes, mime_type: str, filename: str = "",
) -> ExtractResult:
    """Extract plain text + per-page offsets from a binary doc.

    Dispatches on mime_type (preferred) with a filename-extension
    fallback for browsers that upload as application/octet-stream.

    Stateless — no DB or filesystem writes. Caller decides what to do
    with the result (typically: store `text` on
    `bronze_files.metadata.docsage_text` JSONB).
    """
    mt = (mime_type or "").lower()
    fn = (filename or "").lower()

    is_pdf  = mt == "application/pdf" or fn.endswith(".pdf")
    is_docx = (mt in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                      "application/msword"}
               or fn.endswith(".docx"))
    is_image = mt.startswith("image/") or fn.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp"))

    if is_pdf:
        return _extract_pdf(content)
    if is_docx:
        return _extract_docx(content)
    if is_image:
        return ExtractResult(
            text="",
            page_offsets=[0],
            status="unsupported_today",
            error_message=(
                "Ảnh chưa được trích xuất nội dung — chờ Qwen2-VL vision-LLM "
                "adapter ở Phase 2. File vẫn được giữ ở Bronze để DocSage "
                "đọc khi adapter sẵn sàng."
            ),
        )

    return ExtractResult(
        text="",
        page_offsets=[0],
        status="failed",
        error_message=f"Định dạng không hỗ trợ: mime={mime_type!r} filename={filename!r}",
    )


# ─── PDF ────────────────────────────────────────────────────────────


def _extract_pdf(content: bytes) -> ExtractResult:
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
    except Exception as e:
        return ExtractResult(
            text="",
            page_offsets=[0],
            status="failed",
            error_message=f"PDF không đọc được: {type(e).__name__}: {e}",
        )

    if reader.is_encrypted:
        # We don't store passwords; surface a clear message rather than
        # crashing the upload flow.
        try:
            ok = reader.decrypt("") in (1, 2)
        except Exception:
            ok = False
        if not ok:
            return ExtractResult(
                text="",
                page_offsets=[0],
                status="failed",
                error_message="PDF được mã hoá — bỏ password rồi upload lại.",
            )

    pages_text: list[str] = []
    failures = 0
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
            failures += 1
        pages_text.append(t)

    # MinerU Pattern 4 — multi-column reading-order reconstruction.
    # pypdf walks the PDF content stream in raw order which interleaves
    # left/right columns line-by-line for 2-column layouts (VN
    # regulations, financial reports). Em ask pdfplumber for word-level
    # bbox + reorder ONLY the pages it detects as multi-column.
    # Single-column pages keep the pypdf output (zero extra cost). If
    # pdfplumber is unavailable, em skip silently — caller still gets
    # the raw pypdf text + Pattern 2/3 work as before.
    reordered_pages = 0
    reorder_result = extract_reading_order_from_pdf(content)
    if reorder_result is not None:
        for entry in reorder_result:
            if entry.column_count >= 2 and entry.text and entry.page_idx < len(pages_text):
                pages_text[entry.page_idx] = entry.text
        reordered_pages = multi_column_page_count(reorder_result)

    # MinerU pattern 2 — strip repeating header/footer/page-number lines.
    # Pure-Python heuristic, no extra dep. Tracks how much we stripped
    # so callers can see if it kicked in.
    original_total_lines = sum(len(p.splitlines()) for p in pages_text)
    pages_text_clean = strip_repeating_lines(pages_text)
    cleaned_total_lines = sum(len(p.splitlines()) for p in pages_text_clean)
    stripped_count = max(0, original_total_lines - cleaned_total_lines)

    # Rebuild page_offsets from CLEANED text so citation contract stays
    # accurate against the post-strip text.
    page_offsets: list[int] = [0]
    for t in pages_text_clean:
        page_offsets.append(page_offsets[-1] + len(t) + 1)  # +1 = newline separator

    text = "\n".join(pages_text_clean)
    char_count = len(text)
    page_count = len(pages_text_clean)

    # MinerU Pattern 1 — emit TEXT blocks per page.
    blocks: list[Block] = []
    for idx, page_text in enumerate(pages_text_clean):
        block_text = page_text.strip()
        if not block_text:
            continue
        block_start = page_offsets[idx]
        block_end = page_offsets[idx + 1] - 1   # -1 for the joining '\n'
        blocks.append(Block(
            type=BlockType.TEXT,
            page_idx=idx,
            char_start=block_start,
            char_end=block_end,
            text=block_text,
        ))

    # MinerU Pattern 3 — extract tables via pdfplumber, emit TABLE blocks
    # alongside the TEXT blocks. Best-effort: if pdfplumber unavailable
    # or table extraction raises, fall through with no tables (TEXT
    # blocks still cover the page content via pypdf — tables flatten
    # into prose, less ideal but doesn't lose content).
    tables = extract_tables_from_pdf(content)
    for tbl in tables:
        if tbl.page_idx >= len(page_offsets) - 1:
            continue
        # TABLE block carries its markdown rendering as `text` so
        # callers that walk blocks with `text_from_blocks()` get a
        # readable serialization. Raw rows + html live in metadata.
        # char_start/char_end point at the parent page's range — the
        # table is logically located on that page, even though its
        # markdown isn't physically inside the merged `text` string
        # (Pattern 5 will add proper bbox if FE needs it).
        block_start = page_offsets[tbl.page_idx]
        block_end = page_offsets[tbl.page_idx + 1] - 1
        # Pattern 5 (BE foundation 2026-05-19) — pdfplumber gives table
        # bbox via find_tables(); em pass it through. FE bbox highlight
        # (Pattern 5 FE half when restructure resumes) consumes this
        # via Block.bbox. None when pdfplumber couldn't compute.
        table_bbox: Optional[Bbox] = None
        if tbl.bbox is not None:
            table_bbox = Bbox(
                x0=tbl.bbox[0], top=tbl.bbox[1],
                x1=tbl.bbox[2], bottom=tbl.bbox[3],
            )
        blocks.append(Block(
            type=BlockType.TABLE,
            page_idx=tbl.page_idx,
            char_start=block_start,
            char_end=block_end,
            text=tbl.markdown,
            metadata={
                "rows":   tbl.rows,
                "html":   tbl.html,
                "n_rows": tbl.n_rows,
                "n_cols": tbl.n_cols,
            },
            bbox=table_bbox,
        ))

    # If every page returned empty string the PDF is image-only —
    # extract_text() on a scanned PDF returns "" for every page. Mark
    # so the caller can ship to OCR (Phase 2) rather than treat as
    # success-with-no-content.
    if char_count == 0 and page_count > 0:
        return ExtractResult(
            text="",
            page_offsets=page_offsets,
            status="unsupported_today",
            page_count=page_count,
            error_message=(
                "PDF không có text-layer (có thể là PDF scan). "
                "OCR sẽ bật ở Phase 2 — file giữ ở Bronze."
            ),
        )

    if failures > 0:
        return ExtractResult(
            text=text,
            page_offsets=page_offsets,
            status="partial",
            page_count=page_count,
            char_count=char_count,
            blocks=blocks,
            header_footer_lines_stripped=stripped_count,
            multi_column_pages_reordered=reordered_pages,
            error_message=f"{failures}/{page_count} trang lỗi extract; còn lại OK.",
        )

    return ExtractResult(
        text=text,
        page_offsets=page_offsets,
        status="ok",
        page_count=page_count,
        char_count=char_count,
        blocks=blocks,
        header_footer_lines_stripped=stripped_count,
        multi_column_pages_reordered=reordered_pages,
    )


# ─── DOCX ───────────────────────────────────────────────────────────


def _extract_docx(content: bytes) -> ExtractResult:
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as e:
        return ExtractResult(
            text="",
            page_offsets=[0],
            status="failed",
            error_message=f"DOCX không đọc được: {type(e).__name__}: {e}",
        )

    # python-docx has no native "page" concept — the document is a
    # paragraph stream. We treat the whole DOCX as one synthetic page;
    # DocSage's source_segment will be (1, 1) for every row.
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # Include table cells — financial reports + contracts use tables.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                paragraphs.append(row_text)

    text = "\n".join(paragraphs)
    char_count = len(text)

    if not text:
        return ExtractResult(
            text="",
            page_offsets=[0, 0],
            status="failed",
            page_count=1,
            error_message="DOCX không có text nội dung.",
        )

    return ExtractResult(
        text=text,
        page_offsets=[0, char_count + 1],
        status="ok",
        page_count=1,
        char_count=char_count,
    )


# ─── Cache key helper ───────────────────────────────────────────────


def content_fingerprint(content: bytes, *, lib_version: str = "") -> str:
    """Stable cache key for the extraction. Same bytes + same lib version
    = same fingerprint = re-use the cached extraction. Bump
    `lib_version` when pypdf/python-docx upgrades change output."""
    h = hashlib.sha256()
    h.update(content)
    h.update(b"|")
    h.update(lib_version.encode())
    return h.hexdigest()
