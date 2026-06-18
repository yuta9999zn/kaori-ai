"""One-shot extractor for product source docs into plain text for review.

Outputs to docs/product/_extracted/ (gitignored) so we can grep + read with Read tool.
"""
from pathlib import Path
import sys
from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "product"
OUT = ROOT / "docs" / "product" / "_extracted"
OUT.mkdir(exist_ok=True)


def extract_docx(path: Path, out_path: Path) -> int:
    doc = Document(path)
    lines: list[str] = []
    lines.append(f"# Extracted from: {path.name}\n")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            lines.append(f"\n{'#' * (int(level) + 1)} {text}\n")
        else:
            lines.append(text)
    for ti, tbl in enumerate(doc.tables):
        lines.append(f"\n### [Table {ti + 1}]\n")
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return len(lines)


def extract_xlsx(path: Path, out_path: Path) -> int:
    wb = load_workbook(path, data_only=True, read_only=True)
    lines: list[str] = [f"# Extracted from: {path.name}\n"]
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"\n## Sheet: {sheet_name}  (rows={ws.max_row}, cols={ws.max_column})\n")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c).replace("\n", " ").replace("|", "/") for c in row]
            if not any(c.strip() for c in cells):
                continue
            lines.append("| " + " | ".join(cells) + " |")
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return len(lines)


def main() -> int:
    targets = [
        ("TAI_LIEU_YEU_CAU_SAN_PHAM_v5.0.docx", "01_TAI_LIEU_SOURCE.md"),
        ("Kaori_AI_PRD_v5.0.docx", "02_PRD.md"),
        ("Kaori_AI_BRD_v3.0.docx", "03_BRD.md"),
        ("Feature_Tree_Kaori_AI_v3.1.xlsx", "04_FEATURE_TREE.md"),
    ]
    summary: list[tuple[str, int, int]] = []
    for src_name, out_name in targets:
        src_path = SRC / src_name
        out_path = OUT / out_name
        if not src_path.exists():
            print(f"SKIP missing: {src_path}", file=sys.stderr)
            continue
        if src_path.suffix == ".docx":
            n = extract_docx(src_path, out_path)
        else:
            n = extract_xlsx(src_path, out_path)
        size_kb = out_path.stat().st_size // 1024
        summary.append((out_name, n, size_kb))
        print(f"OK  {src_name:40s} -> {out_name}  ({n} lines, {size_kb} KB)")
    print("\nDone. Outputs in:", OUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
